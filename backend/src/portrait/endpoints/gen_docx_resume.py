from datetime import datetime
from io import BytesIO
import json

from django.http import JsonResponse, HttpResponse
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_http_methods

from .common import *
from .ml_utils import generate_general_interpretation_with_ai

# Импорт для DOCX
try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def format_gender(gender_code):
    """Преобразует код пола в читаемый формат."""
    if not gender_code:
        return "Не указан"
    
    gender_map = {
        'м': 'Мужской',
        'М': 'Мужской',
        'ж': 'Женский',
        'Ж': 'Женский'
    }
    
    return gender_map.get(gender_code.lower(), gender_code)


@csrf_exempt
@require_http_methods(["GET"])
def generate_docx_resume(request):
    """
    Генерирует профессиональное резюме в формате DOCX.
    Теперь включает AI-интерпретацию компетенций студента.
    """
    try:
        if not DOCX_AVAILABLE:
            return JsonResponse({
                'status': 'error',
                'message': 'python-docx не установлен. Выполните: pip install python-docx'
            }, status=500)
        
        student_id = request.GET.get('student_id')
        with_ai = request.GET.get('with_ai', 'true').lower() == 'true'  # флаг использования AI
        print(f"📄 Генерация HR-резюме для студента {student_id}, with_ai={with_ai}")
        
        if not student_id:
            return JsonResponse({'status': 'error', 'message': 'student_id обязателен'}, status=400)
        
        # Получаем данные студента
        try:
            participant = Participants.objects.select_related(
                'part_institution', 'part_spec', 'part_form', 'part_edu_level'
            ).get(part_id=student_id)
        except Participants.DoesNotExist:
            return JsonResponse({'status': 'error', 'message': f'Студент {student_id} не найден'}, status=404)
        
        # Получаем последние результаты для компетенций
        latest_result = Results.objects.filter(res_participant=participant).order_by('-res_year', '-res_course_num').first()
        
        # Собираем данные для AI
        competencies_dict = {}
        student_info = {
            'direction': participant.part_spec.spec_name if participant.part_spec else '',
            'course': participant.part_course_num or 1,
            'form': participant.part_form.form_name if participant.part_form else '',
            'level': participant.part_edu_level.edu_level_name if participant.part_edu_level else ''
        }
        
        # Список компетенций в порядке, который используется в ml_utils
        competencies_order = [
            'res_comp_leadership', 'res_comp_communication', 'res_comp_self_development',
            'res_comp_result_orientation', 'res_comp_stress_resistance', 'res_comp_client_focus',
            'res_comp_planning', 'res_comp_info_analysis', 'res_comp_partnership',
            'res_comp_rules_compliance', 'res_comp_emotional_intel', 'res_comp_passive_vocab'
        ]
        
        if latest_result:
            for comp in competencies_order:
                score = getattr(latest_result, comp, None)
                if score and score > 0:
                    competencies_dict[COMP.names[comp]] = score
        
        # Генерация AI-интерпретации, если включено и есть данные
        ai_text = None
        if with_ai and competencies_dict:
            try:
                ai_text = generate_general_interpretation_with_ai(student_info, competencies_dict)
                # Ограничим длину для резюме (не более 3 предложений)
                if ai_text and len(ai_text) > 300:
                    ai_text = ai_text[:300] + '...'
            except Exception as e:
                print(f"AI generation failed: {e}")
                ai_text = None
        
        # Создаём документ
        doc = Document()
        
        # Настройка полей (стандартные для резюме)
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2)
        
        # ============================================================
        # ШАПКА РЕЗЮМЕ
        # ============================================================
        
        # ФИО (крупным жирным шрифтом)
        name_para = doc.add_paragraph()
        name_run = name_para.add_run('ФИО')
        name_run.font.size = Pt(20)
        name_run.font.bold = True
        name_run.font.color.rgb = RGBColor(0, 0, 0)
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_para.space_after = Pt(6)
        
        # Должность
        position_para = doc.add_paragraph()
        position_run = position_para.add_run('Резюме на должность')
        position_run.font.size = Pt(14)
        position_run.font.color.rgb = RGBColor(80, 80, 80)
        position_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        position_para.space_after = Pt(8)
        
        # Контакты
        contact_para = doc.add_paragraph()
        contact_run = contact_para.add_run('Телефон: +7 (ХХХ) ХХХ-ХХ-ХХ\nEmail: email@example.com')
        contact_run.font.size = Pt(11)
        contact_run.font.color.rgb = RGBColor(100, 100, 100)
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_para.space_after = Pt(18)
        
        # ============================================================
        # ЛИЧНАЯ ИНФОРМАЦИЯ
        # ============================================================
        
        heading = doc.add_heading('ЛИЧНАЯ ИНФОРМАЦИЯ', level=1)
        heading.runs[0].font.size = Pt(14)
        heading.runs[0].font.color.rgb = RGBColor(0, 70, 140)
        heading.space_after = Pt(10)
        
        gender = format_gender(participant.part_gender)
        edu_level = participant.part_edu_level.edu_level_name if participant.part_edu_level else 'Не указано'
        
        info_items = [
            f"Место проживания: Место проживания:",
            f"Образование: {edu_level}",
            f"Дата рождения: дд.мм.гггг",
            f"Пол: {gender}"
        ]
        
        for item in info_items:
            p = doc.add_paragraph(item, style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.25)
            p.runs[0].font.size = Pt(11)
            p.space_after = Pt(4)
        
        doc.add_paragraph().space_after = Pt(6)
        
        # ============================================================
        # ОПЫТ РАБОТЫ
        # ============================================================
        
        heading = doc.add_heading('ОПЫТ РАБОТЫ', level=1)
        heading.runs[0].font.size = Pt(14)
        heading.runs[0].font.color.rgb = RGBColor(0, 70, 140)
        heading.space_after = Pt(10)
        
        # Пустые строки для заполнения студентом
        for _ in range(4):
            p = doc.add_paragraph()
            p.add_run('_' * 100).font.color.rgb = RGBColor(200, 200, 200)
            p.space_after = Pt(8)
        
        doc.add_paragraph().space_after = Pt(6)
        
        # ============================================================
        # ОБРАЗОВАНИЕ
        # ============================================================
        
        heading = doc.add_heading('ОБРАЗОВАНИЕ', level=1)
        heading.runs[0].font.size = Pt(14)
        heading.runs[0].font.color.rgb = RGBColor(0, 70, 140)
        heading.space_after = Pt(10)
        
        edu_items = [
            f"Уровень образования: {edu_level}",
            f"Учебное заведение: {participant.part_institution.inst_name if participant.part_institution else 'Не указано'}",
            f"Год окончания: ____________________",
            f"Специальность: {participant.part_spec.spec_name if participant.part_spec else 'Не указано'}",
            f"Форма обучения: {participant.part_form.form_name if participant.part_form else 'Не указана'}"
        ]
        
        for item in edu_items:
            p = doc.add_paragraph(item, style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.25)
            p.runs[0].font.size = Pt(11)
            p.space_after = Pt(4)
        
        doc.add_paragraph().space_after = Pt(6)
        
        # ============================================================
        # ДОПОЛНИТЕЛЬНАЯ ИНФОРМАЦИЯ (с AI-интерпретацией)
        # ============================================================
        
        heading = doc.add_heading('ДОПОЛНИТЕЛЬНАЯ ИНФОРМАЦИЯ', level=1)
        heading.runs[0].font.size = Pt(14)
        heading.runs[0].font.color.rgb = RGBColor(0, 70, 140)
        heading.space_after = Pt(10)
        
        # Если есть AI-текст, добавляем его как выделенный абзац
        if ai_text:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.space_after = Pt(8)
            run = p.add_run("✨ Характеристика по результатам диагностики: ")
            run.font.bold = True
            run.font.size = Pt(11)
            run = p.add_run(ai_text)
            run.font.size = Pt(11)
            run.font.italic = True
            run.font.color.rgb = RGBColor(60, 60, 60)
            doc.add_paragraph().space_after = Pt(6)
        
        # Остальные поля для заполнения (языки, навыки и т.д.)
        additional_info_items = [
            'Иностранные языки: _______________________________________________',
            'Компьютерные навыки: _______________________________________________',
            'Наличие водительских прав (категории): _______________________________________________',
            'Наличие медицинской книжки: _______________________________________________',
            'Занятия в свободное время: _______________________________________________',
            'Личные качества: _______________________________________________'
        ]
        
        for item in additional_info_items:
            p = doc.add_paragraph(item, style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.25)
            p.runs[0].font.size = Pt(11)
            p.space_after = Pt(6)
        
        # Примечание для "Личные качества"
        note_para = doc.add_paragraph()
        note_para.paragraph_format.left_indent = Inches(0.5)
        note_run = note_para.add_run(
            'Примечание: В разделе "Личные качества" укажите ваши профессиональные качества, '
            'такие как ответственность, коммуникабельность, стрессоустойчивость, '
            'целеустремлённость и т.д.'
        )
        note_run.font.size = Pt(9)
        note_run.font.italic = True
        note_run.font.color.rgb = RGBColor(120, 120, 120)
        note_para.space_after = Pt(12)
        
        # ============================================================
        # ФУТЕР
        # ============================================================
        
        footer_para = doc.add_paragraph()
        footer_run = footer_para.add_run(
            f'Дата формирования резюме: {datetime.now().strftime("%d.%m.%Y")}'
        )
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = RGBColor(150, 150, 150)
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # ============================================================
        # СОХРАНЕНИЕ
        # ============================================================
        
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        response = HttpResponse(
            buffer,
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
        filename = f"Resume_Student_{student_id}.docx"
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        
        print(f"✅ HR-резюме сгенерировано: {filename}")
        
        return response
        
    except Exception as e:
        print(f"❌ Ошибка генерации резюме: {str(e)}")
        import traceback
        traceback.print_exc()
        
        return JsonResponse({
            'status': 'error',
            'message': f'Ошибка генерации резюме: {str(e)}'
        }, status=500)