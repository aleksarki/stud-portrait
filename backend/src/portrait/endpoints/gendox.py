from collections import defaultdict
from datetime import datetime
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from .common import *
from .ainterp import *
from ..llmclient import LLM_CLIENT


# ! ================================================== ENDPOINTS =================================================== ! #

@method(GET)
@httpResponse
@csrf_exempt
def generate_docx_resume(request):
    """ Generate professional Docx resume with two-column layout.
    """
    # Получаем параметры из запроса
    student_id = request.GET.get('student_id')
    with_ai = request.GET.get('with_ai', 'true').lower() == 'true'
    
    # Пользовательские данные из конструктора
    phone = request.GET.get('phone', '')
    email = request.GET.get('email', '')
    birth_date = request.GET.get('birth_date', '')
    graduation_year = request.GET.get('graduation_year', '')
    work_experience = request.GET.get('work_experience', '')
    courses = request.GET.get('courses', '')
    languages = request.GET.get('languages', '')
    skills = request.GET.get('skills', '')

    if not student_id:
        raise ResponseError("Missing student_id")

    # Получаем данные студента
    try:
        participant = Participants.objects.get(**{tPART.ID: student_id})
    except Participants.DoesNotExist:
        raise ResponseError(f"Participant {student_id} not found", status=404)

    # Ищем маппинг для получения ФИО
    try:
        mapping = StudentMapping.objects.get(mapping_rsv=participant.part_rsv)
        student_name = mapping.mapping_stud_name
    except StudentMapping.DoesNotExist:
        student_name = participant.part_rsv or f"Участник {student_id}"

    # Последние результаты для компетенций
    latest_result = TestResults.objects                       \
        .filter(**{tRES.PARTICIPANT: participant})        \
        .order_by(desc(tRES.YEAR), desc(tRES.COURSE_NUM)) \
        .first()

    # Данные для ИИ
    competencies_dict = {}
    student_info = {
        'direction': '',
        'course':    participant.part_course_num or 1,
        'form':      '',
        'level':     '',
        'institution': ''
    }

    if latest_result:
        student_info['direction'] = latest_result.res_edu_specialty.edu_spec_name if latest_result.res_edu_specialty else ''
        student_info['form'] = latest_result.res_edu_form.edu_form_name if latest_result.res_edu_form else ''
        student_info['level'] = latest_result.res_edu_level.edu_level_name if latest_result.res_edu_level else ''
        student_info['institution'] = latest_result.res_institution.inst_name if latest_result.res_institution else ''

    # Список компетенций в порядке для отображения
    competencies_order = [
        'res_comp_leadership', 'res_comp_communication', 'res_comp_self_development',
        'res_comp_result_orientation', 'res_comp_stress_resistance', 'res_comp_client_focus',
        'res_comp_planning', 'res_comp_info_analysis', 'res_comp_partnership',
        'res_comp_rules_compliance', 'res_comp_emotional_intel', 'res_comp_passive_vocab'
    ]

    if latest_result:
        for comp in competencies_order:
            score = getattr(latest_result, comp, None)
            if score and score > 0:
                competencies_dict[COMP.names[comp]] = score

    # Если навыки не указаны вручную, генерируем через AI
    if not skills and with_ai and competencies_dict:
        try:
            skills = generate_skills_from_competencies(competencies_dict, student_info)
        except:
            skills = ", ".join([name for name, _ in sorted(competencies_dict.items(), key=lambda x: x[1], reverse=True)[:5]])

    # Генерация AI-интерпретации
    ai_text = None
    if with_ai and competencies_dict:
        try:
            ai_text = generate_general_interpretation_with_ai_for_resume(student_info, competencies_dict)
        except:
            ai_text = None

    # Создаем документ
    doc = Document()

    # Настройка полей
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    # ============================================================
    # ШАПКА РЕЗЮМЕ
    # ============================================================

    # ФИО (крупным жирным шрифтом)
    name_para = doc.add_paragraph()
    name_run = name_para.add_run(student_name)
    name_run.font.size = Pt(22)
    name_run.font.bold = True
    name_run.font.color.rgb = RGBColor(0, 0, 0)
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_para.space_after = Pt(4)

    # Должность/заголовок
    position_para = doc.add_paragraph()
    position_run = position_para.add_run('Резюме')
    position_run.font.size = Pt(14)
    position_run.font.color.rgb = RGBColor(80, 80, 80)
    position_run.font.italic = True
    position_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    position_para.space_after = Pt(12)

    # ============================================================
    # ОСНОВНАЯ ТАБЛИЦА С ДВУМЯ КОЛОНКАМИ
    # ============================================================

    # Создаем таблицу с двумя колонками
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    
    # Настройка ширины колонок (узкая 30%, широкая 70%)
    table.columns[0].width = Cm(5.5)
    table.columns[1].width = Cm(13.5)
    
    # Отключаем границы таблицы
    for row in table.rows:
        for cell in row.cells:
            # Получаем или создаем элемент tcPr
            tcPr = cell._element.tcPr
            if tcPr is None:
                tcPr = OxmlElement('w:tcPr')
                cell._element.append(tcPr)
            
            # Удаляем существующие границы если есть
            borders = tcPr.find(qn('w:tcBorders'))
            if borders is not None:
                tcPr.remove(borders)
            
            # Создаем новые границы без линий
            borders = OxmlElement('w:tcBorders')
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'none')
                border.set(qn('w:sz'), '0')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), 'auto')
                borders.append(border)
            tcPr.append(borders)
    
    # Получаем ячейки
    left_cell = table.cell(0, 0)
    right_cell = table.cell(0, 1)
    
    # Выравнивание по верхнему краю
    left_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    right_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    
    # Отступы внутри ячеек
    left_cell.margin_top = Cm(0.3)
    left_cell.margin_bottom = Cm(0.3)
    left_cell.margin_left = Cm(0.3)
    left_cell.margin_right = Cm(0.3)
    
    right_cell.margin_top = Cm(0.3)
    right_cell.margin_bottom = Cm(0.3)
    right_cell.margin_left = Cm(0.5)
    right_cell.margin_right = Cm(0.3)

    # ============================================================
    # ЛЕВАЯ КОЛОНКА (Узкая)
    # ============================================================

    # --- КОНТАКТЫ ---
    h = left_cell.add_paragraph()
    h_run = h.add_run("КОНТАКТЫ")
    h_run.font.bold = True
    h_run.font.size = Pt(11)
    h_run.font.color.rgb = RGBColor(0, 70, 140)
    h.space_after = Pt(6)
    h.paragraph_format.keep_with_next = True

    if phone:
        p = left_cell.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.add_run(phone).font.size = Pt(9)
    
    if email:
        p = left_cell.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.add_run(email).font.size = Pt(9)
    
    if birth_date:
        p = left_cell.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.add_run(birth_date).font.size = Pt(9)
    
    # Разделитель
    left_cell.add_paragraph().paragraph_format.space_after = Pt(10)

    # --- НАВЫКИ ---
    if skills:
        h = left_cell.add_paragraph()
        h_run = h.add_run("НАВЫКИ")
        h_run.font.bold = True
        h_run.font.size = Pt(11)
        h_run.font.color.rgb = RGBColor(0, 70, 140)
        h.space_after = Pt(6)
        h.paragraph_format.keep_with_next = True
        
        # Разбиваем навыки по запятой
        skills_list = [s.strip() for s in skills.split(",") if s.strip()]
        for skill in skills_list[:10]:
            p = left_cell.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Inches(0.2)
            p.add_run("• ").font.size = Pt(9)
            p.add_run(skill).font.size = Pt(9)
        
        left_cell.add_paragraph().paragraph_format.space_after = Pt(6)

    # --- ЯЗЫКИ ---
    if languages:
        h = left_cell.add_paragraph()
        h_run = h.add_run("ЯЗЫКИ")
        h_run.font.bold = True
        h_run.font.size = Pt(11)
        h_run.font.color.rgb = RGBColor(0, 70, 140)
        h.space_after = Pt(6)
        h.paragraph_format.keep_with_next = True
        
        languages_list = [l.strip() for l in languages.split(",") if l.strip()]
        for lang in languages_list:
            p = left_cell.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Inches(0.2)
            p.add_run("• ").font.size = Pt(9)
            p.add_run(lang).font.size = Pt(9)

    # Добавляем пустое пространство внизу левой колонки
    for _ in range(5):
        left_cell.add_paragraph()

    # ============================================================
    # ПРАВАЯ КОЛОНКА (Широкая)
    # ============================================================

    # --- ОПЫТ РАБОТЫ ---
    if work_experience:
        h = right_cell.add_paragraph()
        h_run = h.add_run("ОПЫТ РАБОТЫ")
        h_run.font.bold = True
        h_run.font.size = Pt(11)
        h_run.font.color.rgb = RGBColor(0, 70, 140)
        h.space_after = Pt(6)
        h.paragraph_format.keep_with_next = True
        
        exp_items = [e.strip() for e in work_experience.split("\n") if e.strip()]
        for exp in exp_items:
            p = right_cell.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.left_indent = Inches(0.2)
            p.add_run("• ").font.size = Pt(10)
            p.add_run(exp).font.size = Pt(10)
        
        right_cell.add_paragraph().paragraph_format.space_after = Pt(8)

    # --- ОБРАЗОВАНИЕ ---
    h = right_cell.add_paragraph()
    h_run = h.add_run("ОБРАЗОВАНИЕ")
    h_run.font.bold = True
    h_run.font.size = Pt(11)
    h_run.font.color.rgb = RGBColor(0, 70, 140)
    h.space_after = Pt(6)
    h.paragraph_format.keep_with_next = True

    edu_level = student_info.get('level', 'Не указано')
    institution = student_info.get('institution', 'Не указано')
    direction = student_info.get('direction', 'Не указано')
    
    edu_items = []
    if institution and institution != 'Не указано':
        edu_items.append(f"Учебное заведение: {institution}")
    if direction and direction != 'Не указано':
        edu_items.append(f"Специальность: {direction}")
    if edu_level and edu_level != 'Не указано':
        edu_items.append(f"Уровень: {edu_level}")
    if graduation_year:
        edu_items.append(f"Год окончания: {graduation_year}")
    
    if not edu_items and graduation_year:
        edu_items.append(f"Год окончания: {graduation_year}")
    
    for item in edu_items:
        p = right_cell.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(0.2)
        p.add_run("• ").font.size = Pt(10)
        p.add_run(item).font.size = Pt(10)
    
    right_cell.add_paragraph().paragraph_format.space_after = Pt(8)

    # --- КУРСЫ ---
    if courses:
        h = right_cell.add_paragraph()
        h_run = h.add_run("КУРСЫ")
        h_run.font.bold = True
        h_run.font.size = Pt(11)
        h_run.font.color.rgb = RGBColor(0, 70, 140)
        h.space_after = Pt(6)
        h.paragraph_format.keep_with_next = True
        
        course_items = [c.strip() for c in courses.split("\n") if c.strip()]
        for course in course_items:
            p = right_cell.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Inches(0.2)
            p.add_run("• ").font.size = Pt(10)
            p.add_run(course).font.size = Pt(10)
        
        right_cell.add_paragraph().paragraph_format.space_after = Pt(8)

    # --- AI-ИНТЕРПРЕТАЦИЯ ---
    if ai_text:
        h = right_cell.add_paragraph()
        h_run = h.add_run("ХАРАКТЕРИСТИКА")
        h_run.font.bold = True
        h_run.font.size = Pt(11)
        h_run.font.color.rgb = RGBColor(0, 70, 140)
        h.space_after = Pt(6)
        h.paragraph_format.keep_with_next = True
        
        p = right_cell.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(0.2)
        p.add_run("• ").font.size = Pt(10)
        
        ai_text_lines = ai_text.split(". ")
        for i, line in enumerate(ai_text_lines):
            if line.strip():
                if i > 0:
                    p = right_cell.add_paragraph()
                    p.paragraph_format.space_after = Pt(2)
                    p.paragraph_format.left_indent = Inches(0.2)
                    p.add_run("  ").font.size = Pt(10)
                p.add_run(line.strip() + ". ").font.size = Pt(10)
                p.add_run().font.italic = True
                p.add_run().font.color.rgb = RGBColor(60, 60, 60)

    # Добавляем пустое пространство внизу правой колонки
    for _ in range(3):
        right_cell.add_paragraph()

    # Возвращаем документ
    return docxResponse(doc, f"Резюме_{student_name}.docx")


@method(GET)
@httpResponse
@csrf_exempt
def generate_geography_report(request):
    """ Generate Word report about the geography of the testing.
    """
    year = request.GET.get('year', '2024/2025')
    available_years = ['2021/2022', '2022/2023', '2023/2024', '2024/2025', '2025/2026']

    if year not in available_years:
        year = '2024/2025'

    doc = Document()

    # Заголовок
    title = doc.add_heading(f'Отчёт о географии тестирования', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Подзаголовок с датой
    date_paragraph = doc.add_paragraph(f'Дата формирования: {datetime.now().strftime("%d.%m.%Y %H:%M")}')
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f'Учебный год: {year}')
    doc.add_paragraph()

    # ========== 1. ГЕОГРАФИЯ ЦЕНТРОВ КОМПЕТЕНЦИЙ ==========
    doc.add_heading('1. География центров компетенций', level=1)

    # Получаем данные по центрам
    results = TestResults.objects        \
        .filter(**{tRES.YEAR: year}) \
        .select_related(tRES.CENTER)

    center_counts = defaultdict(int)
    for result in results:
        if result.res_center:
            center_name = result.res_center.center_name
            center_counts[center_name] += 1

    # Группируем по регионам
    region_stats = defaultdict(int)
    center_region_map = {}

    for center_name, count in center_counts.items():
        region = CENTERS_REGIONS.get(center_name)
        if region:
            region_stats[region] += count
            if region not in center_region_map:
                center_region_map[region] = []
            center_region_map[region].append(center_name)

    # Таблица регионов
    doc.add_heading('Распределение центров по регионам:', level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Заголовки таблицы
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Регион'
    header_cells[1].text = 'Количество центров'
    header_cells[2].text = 'Процент от общего числа'

    total_regions = sum(region_stats.values())

    for region, count in sorted(region_stats.items(), key=lambda x: x[1], reverse=True):
        row_cells = table.add_row().cells
        row_cells[0].text = region
        row_cells[1].text = str(count)
        percent = round(count / total_regions * 100, 1) if total_regions > 0 else 0
        row_cells[2].text = f'{percent}%'

    doc.add_paragraph()

    # Список центров по регионам
    doc.add_heading('Центры компетенций по регионам:', level=2)
    for region, centers in sorted(center_region_map.items(), key=lambda x: x[0]):
        doc.add_paragraph(region)
        for center in centers:
            doc.add_paragraph(center, style='List Bullet')

    doc.add_page_break()

    # ========== 2. СТАТИСТИКА ПО ГОДАМ ==========
    doc.add_heading('2. Динамика по годам', level=1)

    yearly_stats = []
    for yr in available_years:
        yr_results = TestResults.objects.filter(**{tRES.YEAR: yr})
        yr_centers = set()
        for res in yr_results:
            if res.res_center:
                yr_centers.add(res.res_center.center_name)
        yearly_stats.append({
            'year': yr,
            'total_results': yr_results.count(),
            'unique_centers': len(yr_centers)
        })

    # Таблица динамики
    doc.add_heading('Динамика тестирования по годам:', level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'

    header_cells = table.rows[0].cells
    header_cells[0].text = 'Учебный год'
    header_cells[1].text = 'Всего тестирований'
    header_cells[2].text = 'Уникальных центров'
    header_cells[3].text = 'Динамика тестирований'

    prev_total = None
    for stat in yearly_stats:
        row_cells = table.add_row().cells
        row_cells[0].text = stat['year']
        row_cells[1].text = str(stat['total_results'])
        row_cells[2].text = str(stat['unique_centers'])

        if prev_total is not None and prev_total > 0:
            growth = round((stat['total_results'] - prev_total) / prev_total * 100, 1)
            growth_text = f'+{growth}%' if growth > 0 else f'{growth}%'
            row_cells[3].text = growth_text
        else:
            row_cells[3].text = '-'

        prev_total = stat['total_results']

    doc.add_paragraph()

    # ========== 3. ГЕОГРАФИЯ СТУДЕНТОВ ==========
    doc.add_heading('3. География студентов', level=1)

    # Получаем данные о студентах
    students_data = []
    for result in results:
        participant = result.res_participant
        if participant:
            # Ищем маппинг для получения ФИО
            try:
                mapping = StudentMapping.objects.get(mapping_rsv=participant.part_rsv)
                student_name = mapping.mapping_stud_name
            except StudentMapping.DoesNotExist:
                student_name = participant.part_rsv
                
            students_data.append({
                'student_id': participant.part_id,
                'rsv_id': participant.part_rsv,
                'name': student_name,
                'institution': result.res_institution.inst_name if result.res_institution else None,
                'specialty': result.res_edu_specialty.edu_spec_name if result.res_edu_specialty else None,
                'course': participant.part_course_num or result.res_course
            })

    # Группировка по вузам
    institution_stats = defaultdict(int)
    for student in students_data:
        if student['institution']:
            institution_stats[student['institution']] += 1

    doc.add_heading('Распределение студентов по учебным заведениям:', level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'

    header_cells = table.rows[0].cells
    header_cells[0].text = 'Учебное заведение'
    header_cells[1].text = 'Количество студентов'
    header_cells[2].text = 'Процент'

    total_students = len(students_data)
    for inst, count in sorted(institution_stats.items(), key=lambda x: x[1], reverse=True)[:20]:
        row_cells = table.add_row().cells
        row_cells[0].text = inst
        row_cells[1].text = str(count)
        percent = round(count / total_students * 100, 1) if total_students > 0 else 0
        row_cells[2].text = f'{percent}%'

    doc.add_paragraph()

    # ========== 4. РАСПРЕДЕЛЕНИЕ ПО КУРСАМ ==========
    doc.add_heading('4. Распределение по курсам', level=1)

    course_stats = defaultdict(int)
    for student in students_data:
        if student['course']:
            course_stats[student['course']] += 1

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'

    header_cells = table.rows[0].cells
    header_cells[0].text = 'Курс'
    header_cells[1].text = 'Количество студентов'
    header_cells[2].text = 'Процент'

    for course in sorted(course_stats.keys()):
        row_cells = table.add_row().cells
        row_cells[0].text = f'{course} курс'
        row_cells[1].text = str(course_stats[course])
        percent = round(course_stats[course] / total_students * 100, 1) if total_students > 0 else 0
        row_cells[2].text = f'{percent}%'

    doc.add_page_break()

    # ========== 5. РАСПРЕДЕЛЕНИЕ ПО СПЕЦИАЛЬНОСТЯМ ==========
    doc.add_heading('5. Топ-20 направлений подготовки', level=1)

    specialty_stats = defaultdict(int)
    for student in students_data:
        if student['specialty']:
            specialty_stats[student['specialty']] += 1

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'

    header_cells = table.rows[0].cells
    header_cells[0].text = 'Направление подготовки'
    header_cells[1].text = 'Количество студентов'
    header_cells[2].text = 'Процент'

    for specialty, count in sorted(specialty_stats.items(), key=lambda x: x[1], reverse=True)[:20]:
        row_cells = table.add_row().cells
        row_cells[0].text = specialty
        row_cells[1].text = str(count)
        percent = round(count / total_students * 100, 1) if total_students > 0 else 0
        row_cells[2].text = f'{percent}%'

    doc.add_paragraph()

    # ========== 6. ИТОГОВАЯ СТАТИСТИКА ==========
    doc.add_heading('6. Итоговая статистика', level=1)

    stats_paragraph = doc.add_paragraph()
    stats_paragraph.add_run(f'• Всего тестирований за {year}: ').bold = True
    stats_paragraph.add_run(f'{len(results)}\n')

    stats_paragraph.add_run(f'• Всего уникальных центров компетенций: ').bold = True
    stats_paragraph.add_run(f'{len(center_counts)}\n')

    stats_paragraph.add_run(f'• Всего регионов с центрами: ').bold = True
    stats_paragraph.add_run(f'{len(region_stats)}\n')

    stats_paragraph.add_run(f'• Всего студентов: ').bold = True
    stats_paragraph.add_run(f'{total_students}\n')

    stats_paragraph.add_run(f'• Всего учебных заведений: ').bold = True
    stats_paragraph.add_run(f'{len(institution_stats)}\n')

    return docxResponse(doc, f"geography_report_{year}.docx")


# ! ================================================== UTILITIES =================================================== ! #

def generate_skills_from_competencies(competencies_dict, student_info):
    """Generate professional skills list from competencies using AI."""
    if not competencies_dict:
        return ""
    
    # Сортируем компетенции по убыванию
    sorted_comps = sorted(competencies_dict.items(), key=lambda x: x[1], reverse=True)
    
    # Формируем промпт для генерации навыков
    prompt = (
        f"Ты — эксперт по развитию карьеры. На основе баллов развития компетенций студента "
        f"сформируй список из 5-7 ключевых навыков для резюме.\n\n"
        f"Направление подготовки: {student_info.get('direction', 'не указано')}\n"
        f"Курс: {student_info.get('course', 'X')}\n\n"
        f"Результаты диагностики компетенций (шкала 200-800):\n"
    )
    
    prompt += "\n".join([f"- {comp}: {val}" for comp, val in sorted_comps[:8]]) + "\n\n"
    
    prompt += (
        "Инструкция:\n"
        "1. Выбери 5-7 наиболее сильных компетенций студента\n"
        "2. Преобразуй их в конкретные профессиональные навыки\n"
        "3. Навыки должны быть практическими и применимыми в работе\n"
        "4. Учитывай направление подготовки студента\n"
        "5. Используй профессиональную терминологию\n"
        "6. Верни ТОЛЬКО список навыков через запятую, без пояснений\n"
        "7. Навыки должны быть краткими (2-4 слова каждый)\n\n"
        "Навыки:"
    )
    
    try:
        skills_text = LLM_CLIENT.generate(prompt, max_length=500, temperature=0.5, top_p=0.9)
        # Очищаем результат от лишних символов
        skills_text = skills_text.replace('\n', '').replace('"', '').replace("'", "")
        # Разбиваем по запятой и чистим
        skills_list = [s.strip() for s in skills_text.split(",") if s.strip()]
        
        # Если получилось слишком много или слишком мало, корректируем
        if len(skills_list) > 10:
            skills_list = skills_list[:7]
        elif len(skills_list) < 3 and competencies_dict:
            # fallback - берем названия топ-5 компетенций
            skills_list = [name for name, _ in sorted_comps[:5]]
        
        return ", ".join(skills_list)
        
    except Exception as e:
        print(f"[model] Error generating skills: {e}")
        # Fallback - берем топ-5 компетенций
        return ", ".join([name for name, _ in sorted_comps[:5]])


def generate_general_interpretation_with_ai_for_resume(student_info, competencies_dict):
    """Generate a professional resume interpretation using AI."""
    if not competencies_dict:
        return ""
    
    # Сортируем компетенции
    sorted_comps = sorted(competencies_dict.items(), key=lambda x: x[1], reverse=True)
    strong = [name for name, _ in sorted_comps[:3]]
    
    def level(value):
        if value < 399:
            return "начальный уровень"
        if value < 599:
            return "средний уровень"
        return "высокий уровень"
    
    # Формируем промпт
    prompt = (
        f"Ты — карьерный консультант. Напиши краткую характеристику студента для использования в резюме. "
        f"Подчеркни сильные стороны студента, не упоминай слабых сторон. "
        f"Не добавляй никакой информации о внешности, возрасте или личных качествах, не указанных в данных.\n\n"
        f"Курс: {student_info.get('course', 'X')}\n"
        f"Направление: {student_info.get('direction', 'не указано')}\n"
        f"Баллы развития компетенций (200-800):\n"
    )
    
    prompt += "\n".join([f"- {comp}: {val} ({level(val)})" for comp, val in sorted_comps[:8]])
    prompt += f"\n\nСильные стороны (наиболее высокие баллы): {', '.join(strong)}"
    prompt += "\n\nХарактеристика (3-4 предложения, профессиональный стиль):"
    
    try:
        text = LLM_CLIENT.generate(prompt, max_length=800, temperature=0.5, top_p=0.85)
        
        # Очищаем текст
        text = text.strip()
        
        # Проверяем на галлюцинации
        if not text or any(phrase in text.lower() for phrase in ['внешность', 'возраст', 'рост']):
            print("[model] Model generated suspicious content, using fallback")
            raise ValueError("Suspicious content detected")
        
        # Если текст слишком короткий, дополняем
        if len(text.split()) < 20:
            text = f"Студент демонстрирует высокий уровень развития компетенций, особенно в области {', '.join(strong)}. "
        
        return text
        
    except Exception as e:
        print(f"[model] Error generating interpretation: {e}")
        # Fallback
        return f"Студент демонстрирует сильные стороны в области {', '.join(strong)}. Обладает развитыми навыками самоорганизации и достижения целей, готов к профессиональному развитию."
    