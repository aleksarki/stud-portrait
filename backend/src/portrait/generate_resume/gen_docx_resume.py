from django.http import JsonResponse, HttpResponse
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_http_methods
from ..models import Participants
from datetime import datetime
from io import BytesIO

# Импорт для DOCX
try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def format_gender(gender_code):
    """Преобразует код пола в читаемый формат."""
    if not gender_code:
        return "Не указан"
    
    gender_map = {
        'м': 'Мужской',
        'М': 'Мужской',
        'ж': 'Женский',
        'Ж': 'Женский'
    }
    
    return gender_map.get(gender_code.lower(), gender_code)


@csrf_exempt
@require_http_methods(["GET"])
def generate_docx_resume(request):
    """
    Генерирует профессиональное резюме в формате DOCX.
    
    GET /portrait/generate-docx-resume/?student_id=123
    
    Returns:
        DOCX файл для скачивания
    """
    try:
        if not DOCX_AVAILABLE:
            return JsonResponse({
                'status': 'error',
                'message': 'python-docx не установлен. Выполните: pip install python-docx'
            }, status=500)
        
        student_id = request.GET.get('student_id')
        print(f"📄 Генерация HR-резюме для студента {student_id}")
        
        if not student_id:
            return JsonResponse({'status': 'error', 'message': 'student_id обязателен'}, status=400)
        
        # Получаем данные студента
        try:
            participant = Participants.objects.select_related(
                'part_institution', 'part_spec', 'part_form', 'part_edu_level'
            ).get(part_id=student_id)
        except Participants.DoesNotExist:
            return JsonResponse({'status': 'error', 'message': f'Студент {student_id} не найден'}, status=404)
        
        # Создаём документ
        doc = Document()
        
        # Настройка полей (стандартные для резюме)
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2)
        
        # ============================================================
        # ШАПКА РЕЗЮМЕ
        # ============================================================
        
        # ФИО (крупным жирным шрифтом)
        name_para = doc.add_paragraph()
        name_run = name_para.add_run('ФИО')
        name_run.font.size = Pt(20)
        name_run.font.bold = True
        name_run.font.color.rgb = RGBColor(0, 0, 0)
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_para.space_after = Pt(6)
        
        # Должность
        position_para = doc.add_paragraph()
        position_run = position_para.add_run('Резюме на должность')
        position_run.font.size = Pt(14)
        position_run.font.color.rgb = RGBColor(80, 80, 80)
        position_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        position_para.space_after = Pt(8)
        
        # Контакты
        contact_para = doc.add_paragraph()
        contact_run = contact_para.add_run('Телефон: +7 (ХХХ) ХХХ-ХХ-ХХ\nEmail: email@example.com')
        contact_run.font.size = Pt(11)
        contact_run.font.color.rgb = RGBColor(100, 100, 100)
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_para.space_after = Pt(18)
        
        # ============================================================
        # ЛИЧНАЯ ИНФОРМАЦИЯ
        # ============================================================
        
        heading = doc.add_heading('ЛИЧНАЯ ИНФОРМАЦИЯ', level=1)
        heading.runs[0].font.size = Pt(14)
        heading.runs[0].font.color.rgb = RGBColor(0, 70, 140)
        heading.space_after = Pt(10)
        
        # Определяем данные
        gender = format_gender(participant.part_gender)
        edu_level = participant.part_edu_level.edu_level_name if participant.part_edu_level else 'Не указано'
        
        # Личная информация списком
        info_items = [
            f"Место проживания: Место проживания:",
            f"Образование: {edu_level}",
            f"Дата рождения: дд.мм.гггг",
            f"Пол: {gender}"
        ]
        
        for item in info_items:
            p = doc.add_paragraph(item, style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.25)
            p.runs[0].font.size = Pt(11)
            p.space_after = Pt(4)
        
        doc.add_paragraph().space_after = Pt(6)
        
        # ============================================================
        # ОПЫТ РАБОТЫ
        # ============================================================
        
        heading = doc.add_heading('ОПЫТ РАБОТЫ', level=1)
        heading.runs[0].font.size = Pt(14)
        heading.runs[0].font.color.rgb = RGBColor(0, 70, 140)
        heading.space_after = Pt(10)
        
        # Пустые строки для заполнения студентом
        for _ in range(4):
            p = doc.add_paragraph()
            p.add_run('_' * 100).font.color.rgb = RGBColor(200, 200, 200)
            p.space_after = Pt(8)
        
        doc.add_paragraph().space_after = Pt(6)
        
        # ============================================================
        # ОБРАЗОВАНИЕ
        # ============================================================
        
        heading = doc.add_heading('ОБРАЗОВАНИЕ', level=1)
        heading.runs[0].font.size = Pt(14)
        heading.runs[0].font.color.rgb = RGBColor(0, 70, 140)
        heading.space_after = Pt(10)
        
        # Образование списком
        edu_items = [
            f"Уровень образования: {edu_level}",
            f"Учебное заведение: {participant.part_institution.inst_name if participant.part_institution else 'Не указано'}",
            f"Год окончания: ____________________",
            f"Специальность: {participant.part_spec.spec_name if participant.part_spec else 'Не указано'}",
            f"Форма обучения: {participant.part_form.form_name if participant.part_form else 'Не указана'}"
        ]
        
        for item in edu_items:
            p = doc.add_paragraph(item, style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.25)
            p.runs[0].font.size = Pt(11)
            p.space_after = Pt(4)
        
        doc.add_paragraph().space_after = Pt(6)
        
        # ============================================================
        # ДОПОЛНИТЕЛЬНАЯ ИНФОРМАЦИЯ
        # ============================================================
        
        heading = doc.add_heading('ДОПОЛНИТЕЛЬНАЯ ИНФОРМАЦИЯ', level=1)
        heading.runs[0].font.size = Pt(14)
        heading.runs[0].font.color.rgb = RGBColor(0, 70, 140)
        heading.space_after = Pt(10)
        
        # Список пунктов для заполнения
        additional_info_items = [
            'Иностранные языки: _______________________________________________',
            'Компьютерные навыки: _______________________________________________',
            'Наличие водительских прав (категории): _______________________________________________',
            'Наличие медицинской книжки: _______________________________________________',
            'Занятия в свободное время: _______________________________________________',
            'Личные качества: _______________________________________________'
        ]
        
        for item in additional_info_items:
            p = doc.add_paragraph(item, style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.25)
            p.runs[0].font.size = Pt(11)
            p.space_after = Pt(6)
        
        # Примечание для "Личные качества"
        note_para = doc.add_paragraph()
        note_para.paragraph_format.left_indent = Inches(0.5)
        note_run = note_para.add_run(
            'Примечание: В разделе "Личные качества" укажите ваши профессиональные качества, '
            'такие как ответственность, коммуникабельность, стрессоустойчивость, '
            'целеустремлённость и т.д.'
        )
        note_run.font.size = Pt(9)
        note_run.font.italic = True
        note_run.font.color.rgb = RGBColor(120, 120, 120)
        note_para.space_after = Pt(12)
        
        # ============================================================
        # ФУТЕР
        # ============================================================
        
        footer_para = doc.add_paragraph()
        footer_run = footer_para.add_run(
            f'Дата формирования резюме: {datetime.now().strftime("%d.%m.%Y")}'
        )
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = RGBColor(150, 150, 150)
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # ============================================================
        # СОХРАНЕНИЕ
        # ============================================================
        
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        # Возвращаем файл
        response = HttpResponse(
            buffer,
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
        filename = f"Resume_Student_{student_id}.docx"
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        
        print(f"✅ HR-резюме сгенерировано: {filename}")
        
        return response
        
    except Exception as e:
        print(f"❌ Ошибка генерации резюме: {str(e)}")
        import traceback
        traceback.print_exc()
        
        return JsonResponse({
            'status': 'error',
            'message': f'Ошибка генерации резюме: {str(e)}'
        }, status=500)